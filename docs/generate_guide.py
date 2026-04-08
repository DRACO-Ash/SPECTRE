#!/usr/bin/env python3
"""Generate the SIPC Operator Guide as a formatted .docx document.

Follows the Bluestaq Ltd Document Style Guide — A4, Segoe UI, navy/gold
colour scheme, accent bar, metric cards, data tables, and callout boxes.

Usage::

    python docs/generate_guide.py

Outputs ``docs/SIPC_Operator_Guide.docx``.
"""

from __future__ import annotations

import os
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsdecls
from docx.shared import Cm, Emu, Inches, Mm, Pt, RGBColor

# ── Bluestaq colour palette ─────────────────────────────────────────────────
COMMAND_NAVY = RGBColor(0x00, 0x22, 0x44)
ALLIANCE_BLUE = RGBColor(0x00, 0x49, 0x90)
OPS_BLUE = RGBColor(0x00, 0x65, 0x9E)
COMMAND_GOLD = RGBColor(0xC8, 0xA4, 0x15)
BODY_TEXT = RGBColor(0x2C, 0x3E, 0x50)
TABLE_ALT_BG = "F5F7FA"
CALLOUT_BG = "FFF9E6"
GREEN = RGBColor(0x27, 0xAE, 0x60)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

OUTPUT_PATH = Path(__file__).parent / "SIPC_Operator_Guide.docx"


# ═══════════════════════════════════════════════════════════════════════════
#  Style helpers
# ═══════════════════════════════════════════════════════════════════════════


def _setup_document() -> Document:
    """Create a new A4 document with Bluestaq defaults."""
    doc = Document()
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.79)
    section.right_margin = Inches(0.79)

    # Default font.
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Segoe UI"
    font.size = Pt(10)
    font.color.rgb = BODY_TEXT
    style.paragraph_format.space_after = Pt(4)
    return doc


def _add_accent_bar(doc: Document) -> None:
    """Insert a four-segment colour bar in the header."""
    header = doc.sections[0].header
    table = header.add_table(rows=1, cols=4, width=Emu(doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    colours = ["002244", "004990", "00659E", "C8A415"]
    for idx, cell in enumerate(table.rows[0].cells):
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), colours[idx])
        shading.set(qn("w:val"), "clear")
        cell._element.get_or_add_tcPr().append(shading)
        cell.paragraphs[0].text = ""
        # Row height via trHeight.
        cell.width = Emu(int((doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin) / 4))

    # Set row height to thin bar.
    tr = table.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), "120")
    trHeight.set(qn("w:hRule"), "exact")
    trPr.append(trHeight)

    # Remove cell borders for clean appearance.
    for cell in table.rows[0].cells:
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        borders = OxmlElement("w:tcBorders")
        for edge in ("top", "left", "bottom", "right"):
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"), "none")
            el.set(qn("w:sz"), "0")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "auto")
            borders.append(el)
        tcPr.append(borders)


def _add_heading_h1(doc: Document, text: str) -> None:
    """Add a top-level heading — 20pt bold, Command Navy."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = "Segoe UI"
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = COMMAND_NAVY


def _add_eyebrow(doc: Document, text: str) -> None:
    """Add an eyebrow label — 9pt ALL CAPS, Command Gold."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.font.name = "Segoe UI"
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = COMMAND_GOLD


def _add_heading_h2(doc: Document, text: str) -> None:
    """Add a section heading — 16pt bold Navy + 3pt gold rule below."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Segoe UI"
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = COMMAND_NAVY

    # Gold rule below via paragraph border.
    pPr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")  # half-points → 3pt
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "C8A415")
    borders.append(bottom)
    pPr.append(borders)


def _add_heading_h3(doc: Document, text: str) -> None:
    """Add a sub-section heading — 13pt bold Alliance Blue."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Segoe UI"
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = ALLIANCE_BLUE


def _add_body(doc: Document, text: str) -> None:
    """Add a body paragraph — 10pt Segoe UI #2C3E50."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Segoe UI"
    run.font.size = Pt(10)
    run.font.color.rgb = BODY_TEXT


def _add_bullet(doc: Document, text: str, level: int = 0) -> None:
    """Add a bullet point with alternating blue/gold markers."""
    colour = OPS_BLUE if level % 2 == 0 else COMMAND_GOLD
    marker = "\u2022 " if level == 0 else "  \u25E6 "
    indent = "    " * level
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(1.2 * (level + 1))
    marker_run = p.add_run(marker)
    marker_run.font.color.rgb = colour
    marker_run.font.size = Pt(10)
    marker_run.font.bold = True
    text_run = p.add_run(text)
    text_run.font.name = "Segoe UI"
    text_run.font.size = Pt(10)
    text_run.font.color.rgb = BODY_TEXT


def _add_data_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    """Add a data table with navy header row and alternating row shading."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row.
    for idx, hdr in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(hdr)
        run.font.name = "Segoe UI"
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = WHITE
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "002244")
        shading.set(qn("w:val"), "clear")
        cell._element.get_or_add_tcPr().append(shading)

    # Data rows.
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.name = "Segoe UI"
            run.font.size = Pt(9)
            run.font.color.rgb = BODY_TEXT
            if r_idx % 2 == 1:
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), TABLE_ALT_BG)
                shading.set(qn("w:val"), "clear")
                cell._element.get_or_add_tcPr().append(shading)


def _add_callout(doc: Document, text: str) -> None:
    """Add a callout box with left gold border and warm background."""
    # Implemented as a single-cell table for reliable background.
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = "Segoe UI"
    run.font.size = Pt(9.5)
    run.font.color.rgb = BODY_TEXT
    run.font.italic = True

    # Background shading.
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), CALLOUT_BG)
    shading.set(qn("w:val"), "clear")
    cell._element.get_or_add_tcPr().append(shading)

    # Left gold border only.
    tcPr = cell._element.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:color"), "C8A415")
    left.set(qn("w:space"), "0")
    borders.append(left)
    for edge in ("top", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        borders.append(el)
    tcPr.append(borders)


def _add_metric_card(doc: Document, metric: str, label: str) -> None:
    """Add a metric card — navy top border, bold value, grey label."""
    table = doc.add_table(rows=2, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Top border on first cell.
    cell_val = table.rows[0].cells[0]
    cell_val.text = ""
    p = cell_val.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(metric)
    run.font.name = "Segoe UI"
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = COMMAND_NAVY

    tcPr = cell_val._element.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "12")
    top.set(qn("w:color"), "002244")
    top.set(qn("w:space"), "0")
    borders.append(top)
    tcPr.append(borders)

    cell_lbl = table.rows[1].cells[0]
    cell_lbl.text = ""
    p2 = cell_lbl.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(label.upper())
    run2.font.name = "Segoe UI"
    run2.font.size = Pt(8)
    run2.font.bold = True
    run2.font.color.rgb = RGBColor(0x5A, 0x7A, 0x9A)


def _add_footer(doc: Document, title: str = "SIPC Operator Guide") -> None:
    """Add footer with title left and page numbers right."""
    footer = doc.sections[0].footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.clear()

    # Left-aligned title.
    run = p.add_run(f"Bluestaq Ltd \u2014 {title}")
    run.font.name = "Segoe UI"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x5A, 0x7A, 0x9A)

    # Tab to right.
    run_tab = p.add_run("\t\t")

    # Page X of Y field codes.
    run_page = p.add_run("Page ")
    run_page.font.name = "Segoe UI"
    run_page.font.size = Pt(8)
    run_page.font.color.rgb = RGBColor(0x5A, 0x7A, 0x9A)

    fld_page = OxmlElement("w:fldSimple")
    fld_page.set(qn("w:instr"), " PAGE ")
    run_p = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    run_p.append(rPr)
    t = OxmlElement("w:t")
    t.text = "1"
    run_p.append(t)
    fld_page.append(run_p)
    p._p.append(fld_page)

    run_of = p.add_run(" of ")
    run_of.font.name = "Segoe UI"
    run_of.font.size = Pt(8)
    run_of.font.color.rgb = RGBColor(0x5A, 0x7A, 0x9A)

    fld_total = OxmlElement("w:fldSimple")
    fld_total.set(qn("w:instr"), " NUMPAGES ")
    run_n = OxmlElement("w:r")
    rPr2 = OxmlElement("w:rPr")
    run_n.append(rPr2)
    t2 = OxmlElement("w:t")
    t2.text = "1"
    run_n.append(t2)
    fld_total.append(run_n)
    p._p.append(fld_total)


def _add_page_break(doc: Document) -> None:
    """Insert a page break."""
    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════
#  Content sections
# ═══════════════════════════════════════════════════════════════════════════


def _section_01_introduction(doc: Document) -> None:
    """Section 1: Introduction to SIPC."""
    _add_eyebrow(doc, "Section 1")
    _add_heading_h1(doc, "Introduction to SIPC")

    _add_body(doc,
        "The Satellite Intercept Planning Console (SIPC) is a real-time orbital "
        "manoeuvre planning tool designed for space defence operators. It enables "
        "rapid assessment of intercept trajectories, conjunction analysis, and "
        "defensive repositioning of space assets."
    )
    _add_body(doc,
        "SIPC provides a fully self-contained, pure-Python orbital mechanics engine. "
        "Operators can compute classical transfers (Lambert, Hohmann, bi-elliptic), "
        "tactical manoeuvres (phasing, CW relative motion, plane changes, J2 drift "
        "planning, collision avoidance), and decision-support analyses (intent prediction, "
        "intercept envelopes, stability assessment, fingerprinting, formation defence, "
        "orbital terrain mapping, min-time intercept) directly from Two-Line Element "
        "(TLE) data, without requiring any external astrodynamics software."
    )
    _add_callout(doc,
        "Key design principle: SIPC is designed for speed-of-relevance decision-making. "
        "All computations run in-browser with sub-second response times, enabling "
        "operators to evaluate multiple intercept options before a conjunction window closes."
    )


def _section_02_system_overview(doc: Document) -> None:
    """Section 2: System Overview."""
    _add_page_break(doc)
    _add_eyebrow(doc, "Section 2")
    _add_heading_h1(doc, "System Overview")

    _add_heading_h2(doc, "Architecture")
    _add_body(doc,
        "SIPC follows a server-side rendering architecture with HTMX for dynamic "
        "partial updates. The backend is built on FastAPI (Python 3.14+) with Jinja2 "
        "templates. All orbital mechanics computations are performed by the sipc.astro "
        "package — a pure-Python library providing classical transfer solvers (Lambert, "
        "Hohmann, bi-elliptic), tactical manoeuvre solvers (phasing, CW relative motion, "
        "plane change, J2 drift, collision avoidance, evasion), advanced analysis "
        "(GEO drift, NMC safety ellipse, manoeuvre classification, detectability), "
        "decision-support solvers (intent prediction, intercept envelope, stability "
        "analysis, fingerprinting, formation defence, orbital terrain, min-time intercept), "
        "and SGP4-based orbit propagation."
    )

    _add_heading_h2(doc, "Technology Stack")
    _add_data_table(doc,
        ["Component", "Technology", "Purpose"],
        [
            ["Web Framework", "FastAPI + Uvicorn", "Async HTTP server with ASGI"],
            ["Templates", "Jinja2 + HTMX", "Server-side rendering with partial swaps"],
            ["Orbit Propagation", "SGP4 (via sgp4 library)", "TLE-based satellite position/velocity"],
            ["Classical Solvers", "sipc.astro.maneuvers", "Lambert, Hohmann, bi-elliptic transfers"],
            ["Tactical Solvers", "sipc.astro.tactical", "Phasing, CW, plane change, J2 drift, COLA"],
            ["Event Detection", "sipc.astro.events", "Apogee, perigee, node crossings"],
            ["Authentication", "SQLAlchemy + bcrypt", "Operator credential management"],
            ["TLE Source", "UDL (Unified Data Library)", "Authoritative TLE retrieval via HTTPS"],
            ["Charting", "Chart.js 4.x", "Trade-space scatter plots in the browser"],
            ["Design System", "Bluestaq Dark Ops", "Navy/gold operational colour scheme"],
        ],
    )


def _section_03_getting_started(doc: Document) -> None:
    """Section 3: Getting Started."""
    _add_page_break(doc)
    _add_eyebrow(doc, "Section 3")
    _add_heading_h1(doc, "Getting Started")

    _add_heading_h2(doc, "Login")
    _add_body(doc,
        "Navigate to the SIPC URL in your browser. Enter your operator username and "
        "password on the login screen. Credentials are managed by the system administrator "
        "and stored as bcrypt hashes — plaintext passwords are never retained."
    )

    _add_heading_h2(doc, "UDL Credentials")
    _add_body(doc,
        "To fetch authoritative TLE data from the Unified Data Library, connect your "
        "UDL account via the connection chip in the navigation bar. Click the UDL chip, "
        "enter your UDL username and password, and press Connect. A green status dot "
        "confirms a successful connection."
    )
    _add_callout(doc,
        "UDL credentials are held in session memory only and are never persisted to disc. "
        "You will need to re-enter them after each login."
    )

    _add_heading_h2(doc, "UDL Data Mode")
    _add_body(doc,
        "UDL tags every elset record with a data classification mode. The session-wide "
        "data mode can be set from the UDL panel after connecting:"
    )
    _add_data_table(doc,
        ["Mode", "Meaning"],
        [
            ["REAL", "Operational live data (default)"],
            ["SIMULATED", "Synthetic scenario data"],
            ["EXERCISE", "Training exercise data"],
            ["TEST", "Development/test data"],
        ],
    )
    _add_body(doc,
        "Two per-panel data mode overrides are also available:"
    )
    _add_bullet(doc, "Threat Sweep — a filter dropdown in the results panel hides rows whose cached TLE was fetched under a different data mode. This is a display-time filter; no re-fetch occurs.")
    _add_bullet(doc, "Intercept Engine — an 'Intercept Data Mode' dropdown on the intercept form re-fetches the red and blue TLEs from UDL under the selected mode immediately before computing.")

    _add_heading_h2(doc, "TLE Source Provider")
    _add_body(doc,
        "UDL aggregates TLEs from multiple source providers (e.g., 18 SDS, LeoLabs). "
        "When you connect, SIPC automatically discovers available sources for your "
        "account by probing the ISS elset. A 'TLE Source' selector appears in the "
        "UDL panel. Leave it blank to accept UDL's default, or select a specific "
        "provider to pin all subsequent TLE fetches to that source."
    )

    _add_heading_h2(doc, "Setting Scenario Time")
    _add_body(doc,
        "Use the Scenario Time panel to define the analysis window. Enter start and stop "
        "times in ISO 8601 format (e.g., 2026-03-15T00:00:00Z). If no scenario time is "
        "set, SIPC defaults to the current UTC time plus 24 hours."
    )


def _section_04_dashboard(doc: Document) -> None:
    """Section 4: Operator Dashboard Walkthrough."""
    _add_page_break(doc)
    _add_eyebrow(doc, "Section 4")
    _add_heading_h1(doc, "Operator Dashboard Walkthrough")

    _add_body(doc,
        "The main dashboard uses a three-column force grid layout, designed for "
        "rapid situational awareness."
    )

    _add_heading_h2(doc, "Blue Forces (Left Column)")
    _add_bullet(doc, "Lists all friendly blue assets added to the session.")
    _add_bullet(doc, "Each asset displays its object name, NORAD catalogue number, and TLE epoch.")
    _add_bullet(doc, "Add assets by NORAD ID or by pasting a raw TLE.")
    _add_bullet(doc, "Remove assets with the \u00D7 button.")

    _add_heading_h2(doc, "Red Tracks (Centre Column)")
    _add_bullet(doc, "Lists threat objects being tracked.")
    _add_bullet(doc, "Same add/remove workflow as blue assets.")
    _add_bullet(doc, "Red tracks are colour-coded with a red accent border.")

    _add_heading_h2(doc, "Intelligence & Planning (Right Column)")
    _add_bullet(doc, "Manoeuvre search configuration and intercept engine controls.")
    _add_bullet(doc, "Orbital event detection (apogee, perigee, node crossings).")
    _add_bullet(doc, "Intercept results with per-burn ΔV breakdown.")
    _add_bullet(doc, "Trade-space scatter plot (when multiple solutions exist).")
    _add_bullet(doc, "Session activity log with real-time streaming.")


def _section_05_assets(doc: Document) -> None:
    """Section 5: Managing Blue Assets & Red Tracks."""
    _add_page_break(doc)
    _add_eyebrow(doc, "Section 5")
    _add_heading_h1(doc, "Managing Blue Assets & Red Tracks")

    _add_heading_h2(doc, "Adding a Satellite")
    _add_body(doc,
        "There are three modes for adding a satellite to your session:"
    )
    _add_bullet(doc, "UDL Fetch — Enter the NORAD catalogue number and SIPC fetches the latest TLE from UDL automatically. Requires an active UDL connection.")
    _add_bullet(doc, "Catalogue Search — Search the on-orbit catalogue by name or NORAD ID. Select from results to add with a single click.")
    _add_bullet(doc, "Manual TLE — Paste a two-line or three-line TLE directly into the text area. Useful for objects not yet in the public catalogue.")

    _add_heading_h2(doc, "Removing a Satellite")
    _add_body(doc,
        "Click the red \u00D7 button next to any asset or track to remove it from "
        "the session. This action is immediate and does not require confirmation. "
        "Removing a satellite does not affect previously computed intercept results."
    )

    _add_heading_h2(doc, "HRR Watchlist")
    _add_body(doc,
        "The HRR (High Rate Revisit) Watchlist sub-tab in the Assets panel provides "
        "instant access to the Joint Capability Office (JCO) HRR satellite list from "
        "UDL. Click 'Fetch HRR' to load the latest notification (up to 3 days lookback). "
        "The panel shows two tables — Blue HRR and Red HRR — each with:"
    )
    _add_data_table(doc,
        ["Column", "Description"],
        [
            ["Actions", "One-click \u2192 Blue and \u2192 Red ingestion buttons (first column)"],
            ["Name", "Common name from HRR notification"],
            ["SATNO", "NORAD catalogue number"],
            ["Country", "Owner country code"],
            ["Orbit", "Orbit regime (LEO / MEO / GEO / HEO)"],
            ["Rank", "JCO threat rank (0\u20135, lower = higher priority)"],
        ],
    )
    _add_body(doc,
        "Clicking \u2192 Blue fetches the TLE from UDL and adds the satellite as a "
        "blue asset in one action. Clicking \u2192 Red does the same as a red track. "
        "A confirmation badge replaces the button immediately on success. The table "
        "supports column sorting by clicking any header."
    )
    _add_callout(doc,
        "HRR objects are also automatically available as target groups in the "
        "Threat Sweep panel, grouped by side (Blue/Red) and rank."
    )


def _section_06_scenario_time(doc: Document) -> None:
    """Section 6: Setting Scenario Time Window."""
    _add_page_break(doc)
    _add_eyebrow(doc, "Section 6")
    _add_heading_h1(doc, "Setting Scenario Time Window")

    _add_body(doc,
        "The scenario time window defines the temporal bounds for all orbital event "
        "detection and manoeuvre planning. It is set in the Scenario Time panel at "
        "the top of the Blue Forces column."
    )

    _add_data_table(doc,
        ["Field", "Format", "Default"],
        [
            ["Start Time", "ISO 8601 (e.g. 2026-03-15T00:00Z)", "Current UTC time"],
            ["Stop Time", "ISO 8601 (e.g. 2026-03-16T00:00Z)", "Start + 24 hours"],
        ],
    )

    _add_callout(doc,
        "Tip: For conjunction analysis, set the scenario window to span at least "
        "two orbital periods of the target object to capture all relevant geometry."
    )


def _section_07_manoeuvre_theory(doc: Document) -> None:
    """Section 7: Orbital Manoeuvre Theory."""
    _add_page_break(doc)
    _add_eyebrow(doc, "Section 7")
    _add_heading_h1(doc, "Classical Orbital Transfers")

    _add_body(doc,
        "SIPC implements three fundamental orbital transfer methods as its classical "
        "solver set. Each has distinct trade-offs between ΔV cost, transfer time, "
        "and geometric constraints. Understanding these trade-offs is essential for "
        "selecting the optimal intercept strategy."
    )

    # Lambert.
    _add_heading_h2(doc, "Lambert Transfer")
    _add_body(doc,
        "The Lambert problem is the boundary-value problem of two-body orbital "
        "mechanics: given two position vectors and a time of flight, determine the "
        "orbit connecting them. SIPC uses the Izzo algorithm (2015) — a robust, "
        "Householder-iteration solver with guaranteed convergence."
    )
    _add_bullet(doc, "Inputs: departure position, arrival position, time of flight.")
    _add_bullet(doc, "Output: required velocity vectors at departure and arrival.")
    _add_bullet(doc, "The ΔV is the difference between the required velocity and the satellite's current velocity at the departure point.")
    _add_bullet(doc, "Supports both short-way and long-way (multi-revolution) transfers.")
    _add_callout(doc,
        "Lambert transfers are the most flexible method — they work for arbitrary "
        "geometries and arbitrary transfer times. However, they may require higher "
        "ΔV than specialised two-burn transfers when the orbits are nearly coplanar."
    )

    # Hohmann.
    _add_heading_h2(doc, "Hohmann Transfer")
    _add_body(doc,
        "The Hohmann transfer is the minimum-energy two-impulse manoeuvre between "
        "two coplanar circular orbits. It uses an elliptical transfer orbit that is "
        "tangent to both the initial and final orbits."
    )
    _add_bullet(doc, "Burn 1 (departure): Raise/lower apogee to touch the target orbit.")
    _add_bullet(doc, "Burn 2 (arrival): Circularise at the target altitude.")
    _add_bullet(doc, "Transfer time is exactly half the period of the transfer ellipse.")
    _add_body(doc,
        "The vis-viva equation governs the velocity at any point on a Keplerian orbit:"
    )
    _add_body(doc,
        "    v\u00B2 = \u03BC (2/r \u2212 1/a)"
    )
    _add_body(doc,
        "where \u03BC is the gravitational parameter, r is the radial distance, and a "
        "is the semi-major axis."
    )

    # Bi-elliptic.
    _add_heading_h2(doc, "Bi-elliptic Transfer")
    _add_body(doc,
        "The bi-elliptic transfer uses three impulses and two intermediate ellipses. "
        "It is more fuel-efficient than Hohmann when the ratio of the final to initial "
        "orbit radius exceeds approximately 11.94."
    )
    _add_bullet(doc, "Burn 1: Boost to a high intermediate apogee (above both orbits).")
    _add_bullet(doc, "Burn 2: At the intermediate apogee, adjust the perigee to match the target orbit.")
    _add_bullet(doc, "Burn 3: At the new perigee, circularise into the target orbit.")
    _add_callout(doc,
        "The bi-elliptic transfer trades increased transfer time for reduced ΔV. "
        "For orbit ratio r\u2082/r\u2081 > 11.94, it is always cheaper than Hohmann. "
        "Between 11.94 and ~15.58, the savings depend on the intermediate apogee altitude."
    )

    # Comparison table.
    _add_heading_h2(doc, "Method Comparison")
    _add_data_table(doc,
        ["Property", "Lambert", "Hohmann", "Bi-elliptic"],
        [
            ["Number of burns", "1 (single impulse)", "2", "3"],
            ["ΔV efficiency", "Variable", "Optimal for r\u2082/r\u2081 < 11.94", "Optimal for r\u2082/r\u2081 > 11.94"],
            ["Transfer time", "User-specified", "Fixed (half period)", "Longer than Hohmann"],
            ["Geometry constraint", "None (arbitrary)", "Coplanar circular", "Coplanar circular"],
            ["Best use case", "General intercept", "Altitude change", "Large altitude ratio"],
            ["Plane change", "Included in ΔV", "Not included", "Not included"],
        ],
    )


def _section_08_tactical_theory(doc: Document) -> None:
    """Section 8: Tactical Manoeuvre Theory."""
    _add_page_break(doc)
    _add_eyebrow(doc, "Section 8")
    _add_heading_h1(doc, "Tactical Manoeuvres")

    _add_body(doc,
        "Beyond classical orbital transfers, SIPC provides tactical manoeuvres "
        "designed for space control operations. These are computationally lightweight, "
        "operationally focused, and optimised for rapid threat response rather than "
        "textbook optimal transfers."
    )

    # Phasing.
    _add_heading_h2(doc, "Phasing Orbit")
    _add_body(doc,
        "The phasing manoeuvre adjusts the satellite's orbital period so that it "
        "arrives at the target's angular position after a specified number of "
        "revolutions. It is the primary tool for along-track rendezvous when both "
        "objects share a similar orbit."
    )
    _add_bullet(doc, "Burn 1: Lower/raise the orbit to change the period (enter phasing orbit).")
    _add_bullet(doc, "Coast: Complete N revolutions in the phasing orbit, closing the angular gap.")
    _add_bullet(doc, "Burn 2: Return to the original orbit (symmetric with Burn 1).")
    _add_bullet(doc, "More revolutions = lower ΔV but longer wait time.")
    _add_callout(doc,
        "Operational use: stealth rendezvous, delayed intercept, constellation spacing, "
        "passive intercept geometry. The phasing orbit is invisible to observers who "
        "only track the manoeuvre burns — the coast phase appears ballistic."
    )

    # CW Relative Motion.
    _add_heading_h2(doc, "Relative Motion (Hill / Clohessy-Wiltshire)")
    _add_body(doc,
        "The CW equations describe linearised relative motion near a circular "
        "reference orbit. SIPC uses them to compute precise impulses for "
        "controlled proximity operations."
    )
    _add_heading_h3(doc, "Radial Separation")
    _add_bullet(doc, "A single radial impulse produces oscillating radial displacement.")
    _add_bullet(doc, "x(t) = (Δvx / n) sin(nt) — radial position at time t.")
    _add_bullet(doc, "Also creates along-track drift: y(t) = -(2Δvx / n)(1 - cos(nt)).")
    _add_bullet(doc, "Use case: rapid defensive spacing, inspection distance changes, collision avoidance.")
    _add_heading_h3(doc, "Along-Track Drift")
    _add_bullet(doc, "A single along-track impulse produces secular (growing) along-track displacement.")
    _add_bullet(doc, "y(t) = Δvy (4sin(nt) - 3nt) / n — secular term dominates over time.")
    _add_bullet(doc, "Use case: satellite shadowing, constellation phasing, passive intercept.")
    _add_heading_h3(doc, "Combined Manoeuvre")
    _add_bullet(doc, "Solves the full 2×2 CW system for simultaneous radial + along-track targets.")
    _add_bullet(doc, "Operator specifies: 'Create X km radial separation and Y km along-track offset in T minutes.'")
    _add_callout(doc,
        "CW manoeuvres are most accurate when the relative separation is small "
        "compared to the orbital radius (< 100 km). For larger separations, use "
        "Lambert or Hohmann transfers instead."
    )

    # Plane Change.
    _add_heading_h2(doc, "Plane Change")
    _add_body(doc,
        "Many tactical intercepts require orbital plane alignment before a transfer "
        "can be executed. SIPC evaluates two strategies:"
    )
    _add_bullet(doc, "Pure plane change: ΔV = 2v sin(Δi/2). Evaluated at both the equatorial node and apogee — apogee is cheaper for eccentric orbits.")
    _add_bullet(doc, "Combined altitude + plane change: Uses the cosine rule to fold the inclination change into a Hohmann-like transfer. Always cheaper than performing them separately.")
    _add_callout(doc,
        "Key question for operators: 'Is it cheaper to change plane first, or during "
        "the transfer?' SIPC answers this automatically by comparing pure and combined costs."
    )

    # J2 Drift.
    _add_heading_h2(doc, "J2 Drift Planner")
    _add_body(doc,
        "Earth's oblateness (J2) causes the right ascension of the ascending node "
        "(RAAN) to precess secularly. SIPC exploits this for fuel-free orbital "
        "plane alignment."
    )
    _add_body(doc,
        "    dΩ/dt = -3/2 · n · J2 · (R_E / p)² · cos(i)"
    )
    _add_bullet(doc, "Prograde orbits (i < 90°): RAAN drifts westward (negative rate).")
    _add_bullet(doc, "Retrograde orbits (i > 90°): RAAN drifts eastward (positive rate).")
    _add_bullet(doc, "Polar orbits (i = 90°): No RAAN drift — J2 cannot help.")
    _add_bullet(doc, "Different altitudes produce different drift rates → natural convergence.")
    _add_body(doc,
        "SIPC computes the natural convergence time between two satellites and the "
        "small altitude change that would accelerate convergence — trading a tiny ΔV "
        "for days or weeks of waiting time."
    )
    _add_callout(doc,
        "J2 drift is extremely powerful for strategic orbit planning. A 50 km altitude "
        "change can shift the differential RAAN rate by several tenths of a degree per "
        "day — potentially saving weeks of waiting for natural alignment."
    )

    # COLA.
    _add_heading_h2(doc, "Collision Avoidance (COLA)")
    _add_body(doc,
        "When a conjunction is predicted, SIPC computes the minimum ΔV to move the "
        "satellite outside the collision probability envelope. Three strategies are "
        "evaluated simultaneously:"
    )
    _add_bullet(doc, "Radial kick — displace the satellite perpendicular to its orbit plane.")
    _add_bullet(doc, "In-track drift — along-track impulse creates growing separation via differential period.")
    _add_bullet(doc, "Out-of-plane burn — normal impulse shifts the orbit plane slightly.")
    _add_body(doc,
        "SIPC selects the cheapest strategy automatically. The operator sees all three "
        "options and their ΔV costs in the result notes."
    )
    _add_callout(doc,
        "For typical LEO conjunctions with 1+ hour warning, along-track burns are "
        "usually cheapest. For short-notice events (< 30 minutes), radial or "
        "out-of-plane burns may be the only viable option."
    )

    # Evasion.
    _add_heading_h2(doc, "Optimal Defensive Evasion")
    _add_body(doc,
        "The evasion planner extends COLA into a full defensive manoeuvre "
        "generator. Given an incoming threat, a fuel budget, and available "
        "warning time, SIPC evaluates multiple burn strategies across multiple "
        "timing offsets to find the optimal evasion burn."
    )
    _add_bullet(doc, "Respects fuel budget constraints — will not exceed available ΔV.")
    _add_bullet(doc, "Evaluates prograde, normal, and radial burns at 25%, 50%, 75%, and 100% of available warning time.")
    _add_bullet(doc, "Reports remaining fuel after the manoeuvre.")
    _add_bullet(doc, "Selects the strategy that achieves the required miss distance with minimum ΔV.")
    _add_callout(doc,
        "The key difference from COLA: evasion respects operational constraints. "
        "A COLA computation may suggest a burn that exceeds your fuel budget. The "
        "evasion planner will find the best achievable miss distance within your constraints."
    )

    # GEO Drift.
    _add_heading_h2(doc, "GEO Drift Orbit")
    _add_body(doc,
        "For geostationary satellites, SIPC computes east-west longitude "
        "relocation via drift orbits. A small SMA change induces a secular "
        "longitude drift, and a symmetric burn at the target longitude "
        "re-circularises the orbit."
    )
    _add_body(doc,
        "    dλ/dt = -3/2 · n_geo · Δa / a_geo"
    )
    _add_bullet(doc, "Raising SMA by ~1 km → westward drift of ~0.01°/day.")
    _add_bullet(doc, "Two symmetric burns: enter drift orbit + stop at target longitude.")
    _add_bullet(doc, "Also supports graveyard orbit transfers (GEO + 300 km, standard end-of-life disposal).")
    _add_callout(doc,
        "GEO drift is extremely fuel-efficient — relocating 30° in longitude "
        "over 30 days costs only a few m/s of ΔV. This makes it ideal for "
        "strategic repositioning of GEO defence assets."
    )

    # NMC / Safety Ellipse.
    _add_heading_h2(doc, "Natural Motion Circumnavigation (NMC)")
    _add_body(doc,
        "The NMC solver computes relative orbits for proximity operations. "
        "In CW dynamics, bounded (non-drifting) relative motion follows a "
        "2:1 ellipse: the along-track amplitude is always twice the radial "
        "amplitude."
    )
    _add_bullet(doc, "Passive safety: if propulsion fails, the inspector naturally drifts away from the target.")
    _add_bullet(doc, "Safety margin = radial amplitude (minimum distance from target).")
    _add_bullet(doc, "Single impulse to establish the relative orbit (radial + optional cross-track).")
    _add_bullet(doc, "Used by NASA, ESA, and DoD for formation flying and satellite inspection.")
    _add_callout(doc,
        "The passive safety guarantee is critical for high-value assets. If the "
        "inspector's propulsion fails during proximity operations, a properly "
        "designed NMC trajectory ensures it drifts harmlessly away rather than "
        "colliding with the target."
    )

    # Manoeuvre Classification.
    _add_heading_h2(doc, "Manoeuvre Classification Engine")
    _add_body(doc,
        "Given two TLEs of the same satellite at different epochs, SIPC "
        "estimates what manoeuvre occurred between them. This transforms SIPC "
        "from a planning tool into a space intelligence tool."
    )
    _add_bullet(doc, "Compares Keplerian elements (Δa, Δe, Δi, ΔRAAN) between epochs.")
    _add_bullet(doc, "Subtracts expected J2 RAAN drift to isolate manoeuvre-induced changes.")
    _add_bullet(doc, "Classifies: altitude change, plane change, phasing, station-keeping, combined, or unknown.")
    _add_bullet(doc, "Estimates ΔV magnitude and burn direction (prograde, normal, radial, combined).")
    _add_bullet(doc, "Reports confidence level (0–100%) for the classification.")

    # Detectability.
    _add_heading_h2(doc, "Intercept Detectability Metric")
    _add_body(doc,
        "For any computed intercept, SIPC assesses how detectable the "
        "manoeuvre would be by ground-based space surveillance networks. "
        "This enables counter-space analysis and operational security planning."
    )
    _add_bullet(doc, "ΔV categorisation: micro (< 1 m/s), small (1–10 m/s), medium (10–100 m/s), large (> 100 m/s).")
    _add_bullet(doc, "Detection probability based on altitude regime (LEO/MEO/GEO) and ΔV magnitude.")
    _add_bullet(doc, "Estimated time to detection (hours) — how quickly ground tracking would notice.")
    _add_bullet(doc, "Overall observability score (0–100) combining all factors.")
    _add_callout(doc,
        "Micro-manoeuvres (< 1 m/s) in GEO are extremely difficult to detect — "
        "detection probability below 10% and potentially days before orbital "
        "changes become apparent. This is why GEO drift manoeuvres are so "
        "strategically significant."
    )

    # Intent Predict.
    _add_heading_h2(doc, "Intent Predict")
    _add_body(doc,
        "The intent predictor scores adversary intercept intent by correlating "
        "observed orbital behaviour against known attack profiles. It combines "
        "manoeuvre classification output with geometric opportunity analysis "
        "to produce a composite threat score."
    )
    _add_bullet(doc, "Analyses recent manoeuvre history (Δa, Δi, ΔRAAN sequence) for convergent behaviour towards a blue asset.")
    _add_bullet(doc, "Geometric opportunity score: evaluates phase angle closure rate, coplanarity trend, and relative altitude convergence.")
    _add_bullet(doc, "Intent score (0–100): weighted combination of behavioural correlation and geometric opportunity.")
    _add_bullet(doc, "Thresholds: 0–25 nominal, 25–50 ambiguous, 50–75 suspicious, 75–100 hostile.")
    _add_callout(doc,
        "Intent prediction is inherently probabilistic — it identifies patterns "
        "consistent with intercept preparation, not certainty. Always combine with "
        "intelligence context and orbital geometry before escalating."
    )

    # Intercept Envelope.
    _add_heading_h2(doc, "Intercept Envelope")
    _add_body(doc,
        "The intercept envelope solver computes the probabilistic reachability "
        "volume for a satellite within a given time horizon. It sweeps time of "
        "flight values across the vis-viva equation to determine the set of "
        "orbits achievable within a ΔV budget."
    )
    _add_bullet(doc, "Vis-viva TOF sweep: evaluates reachable semi-major axes for ΔV budgets from 0.01 to 5.0 km/s.")
    _add_bullet(doc, "Produces altitude-vs-time contour showing the reachability envelope.")
    _add_bullet(doc, "Identifies which blue assets fall within the adversary's reachable set at each time step.")
    _add_bullet(doc, "Supports both prograde-only and arbitrary-direction burn assumptions.")
    _add_callout(doc,
        "The intercept envelope answers the key operational question: 'Which of "
        "our assets can this threat reach, and how quickly?' Use it for force "
        "protection prioritisation and early warning."
    )

    # Stability Analysis.
    _add_heading_h2(doc, "Stability Analysis")
    _add_body(doc,
        "The stability analyser evaluates the boundedness of relative motion "
        "between two co-orbital objects using the CW drift condition. Bounded "
        "(non-drifting) relative motion requires the initial along-track velocity "
        "offset to satisfy:"
    )
    _add_body(doc,
        "    Δvy₀ = −2n · Δx₀"
    )
    _add_bullet(doc, "Computes the stability score: ratio of actual Δvy₀ to the bounded-motion requirement.")
    _add_bullet(doc, "Score near 1.0 = bounded (formation-stable); score ≫ 1.0 or ≪ 1.0 = drifting apart or converging.")
    _add_bullet(doc, "Reports the secular drift rate in km/orbit if the condition is not satisfied.")
    _add_bullet(doc, "Use case: assess whether an adversary is passively co-orbiting or actively maintaining proximity.")
    _add_callout(doc,
        "A stability score near 1.0 for an adversary satellite is a strong indicator "
        "of deliberate station-keeping relative to your asset — this is a key input "
        "to the intent predictor."
    )

    # Fingerprint.
    _add_heading_h2(doc, "Manoeuvre Fingerprinting")
    _add_body(doc,
        "The fingerprint engine extends manoeuvre classification by matching "
        "observed behaviour against six canonical manoeuvre profiles. Each "
        "profile captures a distinct operational pattern with characteristic "
        "orbital element signatures."
    )
    _add_bullet(doc, "Profile 1 — Station-keeping: periodic small Δa corrections, stable inclination and RAAN.")
    _add_bullet(doc, "Profile 2 — Phasing: periodic Δa oscillations with monotonic phase-angle closure.")
    _add_bullet(doc, "Profile 3 — Plane alignment: progressive Δi or ΔRAAN changes toward a target plane.")
    _add_bullet(doc, "Profile 4 — Intercept approach: combined altitude and plane convergence toward a specific target.")
    _add_bullet(doc, "Profile 5 — Evasion: abrupt manoeuvre followed by divergent trajectory relative to a threat.")
    _add_bullet(doc, "Profile 6 — Disposal: monotonic altitude decrease or GEO graveyard insertion.")
    _add_bullet(doc, "Returns the best-matching profile, correlation score (0–1), and confidence level.")
    _add_callout(doc,
        "Fingerprinting is most effective when multiple TLE epochs are available "
        "(3+ manoeuvre observations). Single-epoch classification reverts to the "
        "standard manoeuvre classification engine."
    )

    # Formation Defence.
    _add_heading_h2(doc, "Formation Defence")
    _add_body(doc,
        "The formation defence solver extends COLA by adding a formation spacing "
        "constraint. When defending a formation of blue assets, collision avoidance "
        "manoeuvres must not only avoid the threat but also maintain the minimum "
        "required spacing between formation members."
    )
    _add_bullet(doc, "Input: formation member positions, minimum inter-satellite spacing, and the conjunction threat.")
    _add_bullet(doc, "Evaluates COLA manoeuvres that maintain formation geometry within tolerance.")
    _add_bullet(doc, "If no single-burn solution preserves formation spacing, recommends coordinated multi-satellite burns.")
    _add_bullet(doc, "Reports post-manoeuvre formation integrity score and any spacing violations.")
    _add_callout(doc,
        "Formation defence is critical for distributed space architectures where "
        "losing formation geometry may be as costly as the conjunction itself. "
        "Always verify formation integrity after any COLA manoeuvre."
    )

    # Orbital Terrain.
    _add_heading_h2(doc, "Orbital Terrain Mapping")
    _add_body(doc,
        "The orbital terrain mapper produces an altitude-vs-inclination risk map "
        "characterising the operational environment. It combines three risk layers "
        "to produce a composite terrain score for any orbital regime."
    )
    _add_bullet(doc, "Debris density: spatial density of tracked debris objects per altitude/inclination bin.")
    _add_bullet(doc, "Congestion: number of active satellites and manoeuvring objects in each bin.")
    _add_bullet(doc, "Radiation: trapped proton and electron flux from the AP-8/AE-8 radiation belt models.")
    _add_bullet(doc, "Composite terrain score (0–100): weighted sum of debris, congestion, and radiation risk.")
    _add_bullet(doc, "Use case: route planning for orbital transfers, identifying safe parking orbits, risk-aware manoeuvre selection.")
    _add_callout(doc,
        "The 800–1000 km altitude band at high inclination (sun-synchronous) "
        "typically scores highest for debris risk due to historical fragmentation "
        "events. GEO scores highest for congestion. MEO radiation belts peak "
        "near 20,000 km altitude."
    )

    # Min-Time Intercept.
    _add_heading_h2(doc, "Min-Time Intercept")
    _add_body(doc,
        "The min-time intercept solver finds the shortest time of flight that "
        "achieves intercept within a given ΔV budget. It uses a binary search "
        "on TOF, evaluating Lambert solutions at each step to find the minimum "
        "transfer time that satisfies the fuel constraint."
    )
    _add_bullet(doc, "Input: blue asset, red track, maximum ΔV budget (km/s).")
    _add_bullet(doc, "Binary search converges on TOF to within 1-second precision.")
    _add_bullet(doc, "Reports the minimum achievable TOF, required ΔV, and the corresponding Lambert solution.")
    _add_bullet(doc, "If no solution exists within the ΔV budget, reports the minimum ΔV required and the corresponding TOF.")
    _add_callout(doc,
        "Min-time intercept answers the urgent operational question: 'How fast can "
        "we get there with the fuel we have?' Use it for time-critical threat "
        "response where speed matters more than fuel efficiency."
    )

    # Extended comparison table.
    _add_heading_h2(doc, "Complete Method Comparison")
    _add_data_table(doc,
        ["Method", "Burns", "Best Use Case", "ΔV Character"],
        [
            ["Lambert", "1–2", "General intercept, arbitrary geometry", "Variable"],
            ["Hohmann", "2", "Altitude change, coplanar circular", "Minimum energy"],
            ["Bi-elliptic", "3", "Large altitude ratio (r₂/r₁ > 11.94)", "Lower than Hohmann for large ratios"],
            ["Phasing", "2", "Along-track rendezvous, delayed intercept", "Very low (period adjustment)"],
            ["CW Radial", "1", "Defensive spacing, inspection geometry", "Very low (metres/s)"],
            ["CW Drift", "1", "Shadowing, constellation spacing", "Very low (metres/s)"],
            ["Plane Change", "1–2", "Inclination alignment", "High (proportional to Δi)"],
            ["J2 Drift", "0–1", "RAAN alignment, strategic planning", "Near-zero (natural precession)"],
            ["COLA", "1", "Conjunction avoidance", "Minimum for required miss distance"],
            ["Evasion", "1", "Defensive manoeuvre under fuel constraints", "Budget-constrained minimum"],
            ["GEO Drift", "2", "GEO longitude relocation", "Very low (mm/s to m/s)"],
            ["NMC", "1", "Proximity ops, formation flying", "Very low (relative orbit establishment)"],
            ["Manoeuvre Detect", "0", "Space intelligence, behaviour analysis", "N/A (analytical)"],
            ["Detectability", "1–2", "Counter-space analysis, OPSEC", "Same as underlying intercept"],
            ["Intent Predict", "0", "Adversary intent scoring, threat assessment", "N/A (analytical)"],
            ["Intercept Envelope", "0", "Reachability analysis, force protection", "Budget-swept (vis-viva)"],
            ["Stability Analysis", "0", "Relative motion boundedness assessment", "N/A (analytical)"],
            ["Fingerprint", "0", "Behavioural classification against 6 profiles", "N/A (analytical)"],
            ["Formation Defence", "1–N", "COLA with formation spacing constraint", "Minimum for miss + spacing"],
            ["Orbital Terrain", "0", "Risk mapping: debris, congestion, radiation", "N/A (analytical)"],
            ["Min-Time Intercept", "1–2", "Fastest intercept within ΔV budget", "Budget-constrained minimum TOF"],
        ],
    )


def _section_09_threat_sweep(doc: Document) -> None:
    """Section 9: Threat Sweep."""
    _add_page_break(doc)
    _add_eyebrow(doc, "Section 9")
    _add_heading_h1(doc, "Threat Sweep")

    _add_body(doc,
        "The Threat Sweep performs a rapid batch intercept assessment of all HRR "
        "objects in a chosen target group against your primary blue asset. It runs "
        "a Hohmann screening pass across five orbital epochs for each target, then "
        "refines the top-five results with Lambert solvers."
    )

    _add_heading_h2(doc, "Operator Workflow")
    _add_bullet(doc, "1. Connect to UDL and fetch the HRR watchlist (Assets \u2192 HRR tab).")
    _add_bullet(doc, "2. Open the Threat Sweep hero tab.")
    _add_bullet(doc, "3. Select a Target Group from the dropdown (e.g. 'Red HRR \u2013 Rank 1').")
    _add_bullet(doc, "4. Click Fetch TLEs \u2014 SIPC retrieves TLEs for all objects in the group and displays a readiness badge (N/M TLEs ready).")
    _add_bullet(doc, "5. Click Sweep Targets \u2014 the batch assessment runs and the ranked results table appears.")
    _add_bullet(doc, "6. Optionally click Refine on any row for a full Lambert refinement of that target.")

    _add_heading_h2(doc, "Sweep Algorithm")
    _add_body(doc,
        "For each target in the selected group, SIPC computes Hohmann \u0394V at five "
        "canonical epochs: T+0 (now), apogee passage, perigee passage, ascending node, "
        "and descending node. All N\u00D75 entries are ranked by \u0394V; the five unique "
        "lowest-cost targets are then refined with Lambert solvers for precision \u0394V "
        "and miss distance."
    )
    _add_data_table(doc,
        ["Result Column", "Description"],
        [
            ["Target", "Satellite name and NORAD ID"],
            ["Orbit", "Orbit regime compatibility badge"],
            ["\u0394V (km/s)", "Minimum Lambert transfer cost (refined)"],
            ["Miss Distance (km)", "Closest approach distance at intercept epoch"],
            ["Epoch", "Optimal manoeuvre start epoch"],
            ["Method", "Screening method (Hohmann/Lambert)"],
        ],
    )

    _add_heading_h2(doc, "Data Mode Filter")
    _add_body(doc,
        "A 'Data Mode' dropdown above the results table filters displayed rows by the "
        "UDL data mode tag of the cached TLE (REAL / SIMULATED / EXERCISE / TEST). "
        "Select 'All' to show every target regardless of data mode. This is a "
        "display-time filter — no re-fetch occurs."
    )

    _add_callout(doc,
        "Orbit regime compatibility is enforced: a LEO blue asset will only sweep "
        "LEO targets; a GEO blue asset will only sweep GEO targets. Incompatible "
        "targets are excluded with an explanatory note."
    )


def _section_10_intercept_calculations(doc: Document) -> None:
    """Section 10: Running Intercept Calculations."""
    _add_page_break(doc)
    _add_eyebrow(doc, "Section 10")
    _add_heading_h1(doc, "Running Intercept Calculations")

    _add_heading_h2(doc, "Method Selection")
    _add_body(doc,
        "Select the intercept method from the dropdown in the Intercept Engine panel. "
        "Methods are grouped into two categories:"
    )
    _add_bullet(doc, "Classical Transfers: Lambert, Hohmann, Bi-elliptic, Proximity, Rendezvous.")
    _add_bullet(doc, "Tactical Manoeuvres: Phasing, CW Radial, CW Drift, Plane Change, J2 Drift, COLA, Evasion.")
    _add_bullet(doc, "Advanced Analysis: GEO Drift, NMC (safety ellipse), Manoeuvre Detect, Detectability.")
    _add_bullet(doc, "Decision Support: Intent Predict, Intercept Envelope, Stability Analysis, Fingerprint, Formation Defence, Orbital Terrain, Min-Time Intercept.")
    _add_body(doc,
        "The form adapts to show relevant fields for each method — time of flight "
        "is shown for Lambert and CW methods; target distance is shown for CW, "
        "COLA, and Proximity methods."
    )

    _add_heading_h2(doc, "Parameters")
    _add_data_table(doc,
        ["Parameter", "Description", "Default"],
        [
            ["Red Satellite", "Target threat track to intercept", "(select from red tracks)"],
            ["Blue Satellite", "Friendly asset performing the manoeuvre", "(select from blue assets)"],
            ["Manoeuvre Start", "UTC epoch for the first burn", "Current UTC time"],
            ["Coast Hours", "Coast duration before first burn (hours)", "1.0"],
            ["Intercept Hours", "Time of flight for Lambert solver (hours)", "6.0"],
            ["Max ΔV", "Upper bound on acceptable ΔV (km/s)", "3.0"],
            ["Intercept Data Mode", "Re-fetch red/blue TLEs from UDL under this data mode before computing. Leave blank to use session TLEs.", "Session default (REAL)"],
        ],
    )

    _add_heading_h2(doc, "Interpreting the Burn Table")
    _add_body(doc,
        "The result displays a per-burn breakdown table with the following columns:"
    )
    _add_bullet(doc, "Burn # — sequential burn index (1-based).")
    _add_bullet(doc, "Segment — descriptive name of the manoeuvre segment.")
    _add_bullet(doc, "Burn Epoch — UTC time of the impulsive burn.")
    _add_bullet(doc, "ΔV km/s — total delta-V magnitude for this burn.")
    _add_bullet(doc, "Prograde — velocity-direction component (VNB frame).")
    _add_bullet(doc, "Normal — orbit-normal component (VNB frame).")
    _add_bullet(doc, "Radial — co-normal/radial component (VNB frame).")

    _add_heading_h2(doc, "VNB Reference Frame")
    _add_body(doc,
        "SIPC decomposes all ΔV vectors into the VNB (Velocity-Normal-Binormal) frame, "
        "also known as the local orbital frame. This is the standard frame used by "
        "spacecraft operators for manoeuvre planning:"
    )
    _add_bullet(doc, "V (Prograde): Along the velocity vector. Positive prograde raises the orbit on the opposite side.")
    _add_bullet(doc, "N (Normal): Perpendicular to the orbital plane (angular momentum direction). Used for plane changes.")
    _add_bullet(doc, "B (Radial/Binormal): Completes the right-hand triad. Points generally away from the central body.")


def _section_11_trade_space(doc: Document) -> None:
    """Section 11: Trade-Space Analysis."""
    _add_page_break(doc)
    _add_eyebrow(doc, "Section 11")
    _add_heading_h1(doc, "Trade-Space Analysis")

    _add_body(doc,
        "The trade-space scatter plot visualises all intercept solutions computed "
        "during the current session. It plots ΔV (km/s) against transfer time (minutes), "
        "with points colour-coded by method:"
    )
    _add_bullet(doc, "Lambert / Proximity — Ops Blue (#00659E)")
    _add_bullet(doc, "Hohmann / Rendezvous — Command Gold (#C8A415)")
    _add_bullet(doc, "Bi-elliptic — Green (#27AE60)")
    _add_bullet(doc, "Tactical manoeuvres (Phasing, CW, Plane Change, J2 Drift, COLA) — distinct colours per method")

    _add_heading_h2(doc, "Using the Trade-Space Plot")
    _add_body(doc,
        "Run intercept calculations with different methods and parameters. After "
        "the second solution, the trade-space chart appears automatically below the "
        "burn table. Hover over any point to see the method, ΔV, transfer time, and "
        "miss distance."
    )
    _add_body(doc,
        "Look for the Pareto front — solutions in the lower-left region of the chart "
        "offer the best combination of low ΔV and short transfer time. Solutions "
        "closer to the axes represent trade-offs: low ΔV with longer transfers, or "
        "fast transfers with higher ΔV."
    )

    _add_heading_h2(doc, "Clearing History")
    _add_body(doc,
        "Click the 'Clear History' button on the chart to reset all accumulated "
        "solutions and start a fresh trade-space analysis. This also clears the "
        "current intercept result display."
    )


def _section_12_scenario_planning(doc: Document) -> None:
    """Section 12: Scenario Planning."""
    _add_page_break(doc)
    _add_eyebrow(doc, "Section 12")
    _add_heading_h1(doc, "Scenario Planning")

    _add_heading_h2(doc, "Access Windows")
    _add_body(doc,
        "SIPC detects orbital events — apogee, perigee, ascending node, and "
        "descending node — for both red and blue satellites within the scenario "
        "time window. These events define natural manoeuvre opportunities."
    )
    _add_bullet(doc, "Apogee burns are efficient for lowering the perigee.")
    _add_bullet(doc, "Perigee burns are efficient for raising the apogee.")
    _add_bullet(doc, "Node crossings indicate plane-change opportunities.")

    _add_heading_h2(doc, "Conjunction Detection")
    _add_body(doc,
        "When both a red track and a blue asset are loaded, SIPC can identify "
        "close-approach windows by propagating both TLEs through the scenario "
        "period. Events are displayed as clickable badges that auto-populate "
        "the manoeuvre start time field."
    )
    _add_callout(doc,
        "Operator tip: Use orbital events to identify the optimal manoeuvre epoch, "
        "then run Lambert solvers at those epochs to minimise ΔV."
    )


def _section_13_scenarios(doc: Document) -> None:
    """Section 13: Operator Scenarios."""
    _add_page_break(doc)
    _add_eyebrow(doc, "Section 13")
    _add_heading_h1(doc, "Operator Scenarios")

    # Scenario 1.
    _add_heading_h2(doc, "Scenario 1: Quick-Look Conjunction Assessment")
    _add_body(doc, "Objective: Rapidly assess whether a conjunction event requires a manoeuvre.")
    _add_bullet(doc, "Add the blue asset by NORAD ID (e.g., 25544 for ISS).")
    _add_bullet(doc, "Add the red track by NORAD ID.")
    _add_bullet(doc, "Set the scenario window to cover the predicted conjunction epoch \u00B1 12 hours.")
    _add_bullet(doc, "Run Orbital Events to identify the closest approach geometry.")
    _add_bullet(doc, "Run a Lambert intercept at the conjunction epoch to determine the ΔV required for a collision avoidance manoeuvre.")
    _add_bullet(doc, "If ΔV < 0.01 km/s, a small adjustment is sufficient. If ΔV > 0.1 km/s, consider alternative timing.")

    # Scenario 2.
    _add_heading_h2(doc, "Scenario 2: Defensive Repositioning of Blue Asset")
    _add_body(doc, "Objective: Move a blue asset to a safer orbit to avoid a persistent threat.")
    _add_bullet(doc, "Add the blue asset and identify its current orbital elements from the TLE.")
    _add_bullet(doc, "Determine the desired target altitude (e.g., raise by 50 km).")
    _add_bullet(doc, "Run a Hohmann transfer to compute the two-burn sequence for the altitude change.")
    _add_bullet(doc, "Review the burn table: verify the total ΔV is within the satellite's remaining fuel budget.")
    _add_bullet(doc, "Note the transfer time — if too long, consider a Lambert solution with a shorter time of flight.")

    # Scenario 3.
    _add_heading_h2(doc, "Scenario 3: Multi-Method Intercept Comparison")
    _add_body(doc, "Objective: Compare Lambert, Hohmann, and bi-elliptic solutions for the same target pair.")
    _add_bullet(doc, "Add one blue asset and one red track.")
    _add_bullet(doc, "Run Lambert with 6-hour time of flight.")
    _add_bullet(doc, "Run Hohmann for the same satellite pair.")
    _add_bullet(doc, "Run Bi-elliptic for the same pair.")
    _add_bullet(doc, "The trade-space chart appears after the second run. Compare all three points.")
    _add_bullet(doc, "Select the solution on the Pareto front that best balances ΔV and transfer time for your operational constraints.")

    # Scenario 4.
    _add_heading_h2(doc, "Scenario 4: Historical TLE Epoch Analysis")
    _add_body(doc, "Objective: Analyse how intercept costs change with different TLE epochs.")
    _add_bullet(doc, "Add the red track using the most recent TLE from UDL.")
    _add_bullet(doc, "Run a Lambert intercept and note the ΔV and miss distance.")
    _add_bullet(doc, "Remove the red track and re-add with an older TLE (e.g., from a previous day).")
    _add_bullet(doc, "Run the same Lambert intercept and compare results.")
    _add_bullet(doc, "Larger miss distances with older TLEs indicate orbital prediction degradation. Use this to assess TLE freshness requirements for your mission.")

    # Scenario 5.
    _add_heading_h2(doc, "Scenario 5: Full Mission Planning Workflow")
    _add_body(doc, "Objective: Plan a complete intercept mission from initial assessment through solution selection.")
    _add_bullet(doc, "Log in and connect to UDL.")
    _add_bullet(doc, "Set the scenario window to the mission planning period.")
    _add_bullet(doc, "Add all relevant blue assets and red tracks.")
    _add_bullet(doc, "Run Orbital Events for each red/blue pair to identify manoeuvre windows.")
    _add_bullet(doc, "For each promising window, run Lambert, Hohmann, and Bi-elliptic solutions.")
    _add_bullet(doc, "Use the trade-space plot to identify the optimal solution.")
    _add_bullet(doc, "Record the selected burn table (epoch, ΔV components) for uplink to the spacecraft.")
    _add_bullet(doc, "Clear history and repeat for the next target pair if required.")

    # Scenario 6.
    _add_heading_h2(doc, "Scenario 6: CW Defensive Spacing")
    _add_body(doc, "Objective: Create a safe standoff distance from a threat using CW proximity manoeuvres.")
    _add_bullet(doc, "Add the blue asset and the red track (both in similar orbits).")
    _add_bullet(doc, "Select CW Radial Separation with a target distance of 5 km and a time of 30 minutes.")
    _add_bullet(doc, "Review the single-impulse ΔV — typically metres per second for short-range proximity operations.")
    _add_bullet(doc, "If along-track displacement is also needed, switch to CW Along-Track Drift and specify the desired drift distance.")
    _add_bullet(doc, "Compare both CW solutions in the trade-space plot — radial manoeuvres are faster but oscillatory; along-track drift is secular (growing).")

    # Scenario 7.
    _add_heading_h2(doc, "Scenario 7: COLA Conjunction Response")
    _add_body(doc, "Objective: Respond to a predicted conjunction by computing the cheapest avoidance manoeuvre.")
    _add_bullet(doc, "Add the blue asset and the red threat object.")
    _add_bullet(doc, "Run Orbital Events to confirm the conjunction geometry and timing.")
    _add_bullet(doc, "Select COLA and set the required miss distance (e.g., 5 km).")
    _add_bullet(doc, "SIPC evaluates all three strategies (radial, in-track, out-of-plane) and recommends the cheapest.")
    _add_bullet(doc, "Review the result notes to see all three strategy costs and the selected option.")
    _add_bullet(doc, "If time permits (> 1 hour to conjunction), the in-track burn is usually cheapest. For urgent conjunctions, radial or cross-track may be the only option.")

    # Scenario 8.
    _add_heading_h2(doc, "Scenario 8: GEO Longitude Relocation")
    _add_body(doc, "Objective: Relocate a GEO satellite 15° eastward over 15 days.")
    _add_bullet(doc, "Add the GEO satellite as a blue asset.")
    _add_bullet(doc, "Add a reference satellite at the target longitude as a red track (or use a GEO satellite already at the desired slot).")
    _add_bullet(doc, "Select GEO Drift and set coast hours to represent the desired drift duration in days.")
    _add_bullet(doc, "Review the two-burn solution: enter drift orbit (small prograde burn) and stop at target longitude (symmetric retrograde burn).")
    _add_bullet(doc, "Verify total ΔV is within the satellite's remaining fuel budget — typical GEO relocations cost only a few m/s.")

    # Scenario 9.
    _add_heading_h2(doc, "Scenario 9: Threat Behaviour Analysis")
    _add_body(doc, "Objective: Classify an observed manoeuvre by an adversary satellite.")
    _add_bullet(doc, "Obtain two TLEs for the same satellite at different epochs (before and after suspected manoeuvre).")
    _add_bullet(doc, "Add the earlier TLE as a red track and the later TLE as a blue asset.")
    _add_bullet(doc, "Select Manoeuvre Detect — SIPC compares the orbital elements and classifies the manoeuvre.")
    _add_bullet(doc, "Review: manoeuvre type (altitude change, plane change, phasing, etc.), estimated ΔV, burn direction, and confidence level.")
    _add_bullet(doc, "Use the classification to assess intent — altitude changes suggest repositioning, plane changes suggest target tracking.")

    # Scenario 10.
    _add_heading_h2(doc, "Scenario 10: Adversary Intent Assessment with Formation Defence")
    _add_body(doc, "Objective: Assess adversary intent toward a blue formation and compute a coordinated defensive response.")
    _add_bullet(doc, "Add all formation members as blue assets and the suspected adversary as a red track.")
    _add_bullet(doc, "Run Intent Predict to score the adversary's intercept intent against each formation member.")
    _add_bullet(doc, "Run Intercept Envelope to determine which formation members fall within the adversary's reachable set.")
    _add_bullet(doc, "Run Stability Analysis on the adversary-to-nearest-blue pair to assess whether the adversary is actively station-keeping relative to the formation.")
    _add_bullet(doc, "Run Fingerprint on the adversary's recent TLE history to classify its behavioural profile.")
    _add_bullet(doc, "If intent score exceeds 50, run Formation Defence to compute a coordinated COLA manoeuvre that preserves formation spacing.")
    _add_bullet(doc, "Review the Orbital Terrain map to identify safe retreat orbits with low debris and congestion risk.")
    _add_bullet(doc, "If immediate response is required, run Min-Time Intercept to find the fastest achievable defensive repositioning within the fuel budget.")


def _section_14_glossary(doc: Document) -> None:
    """Section 14: Glossary of Terms."""
    _add_page_break(doc)
    _add_eyebrow(doc, "Section 14")
    _add_heading_h1(doc, "Glossary of Terms")

    terms = [
        ["Apogee", "The point in an orbit furthest from the central body."],
        ["Bi-elliptic Transfer", "A three-burn orbital transfer using two intermediate ellipses."],
        ["COLA", "Collision Avoidance — minimum-ΔV manoeuvre to avoid a predicted conjunction."],
        ["Conjunction", "A close approach between two space objects."],
        ["Coast", "A ballistic (unpowered) phase between manoeuvres."],
        ["CW Equations", "Clohessy-Wiltshire (Hill) equations for linearised relative motion near a circular orbit."],
        ["ΔV (Delta-V)", "Change in velocity; the fundamental currency of orbital manoeuvres."],
        ["Detectability", "The likelihood that a manoeuvre will be observed by space surveillance networks."],
        ["Drift Orbit", "An orbit with a slightly different SMA than GEO, causing east-west longitude drift."],
        ["Fingerprinting", "Behavioural classification of a satellite's manoeuvre history against canonical operational profiles."],
        ["Formation Defence", "Collision avoidance manoeuvre planning that maintains minimum inter-satellite spacing within a formation."],
        ["Epoch", "A specific point in time, usually UTC."],
        ["Evasion", "A defensive manoeuvre to avoid an incoming threat while respecting fuel constraints."],
        ["GEO", "Geostationary Earth Orbit — circular orbit at 42,164 km radius with 24-hour period."],
        ["Graveyard Orbit", "Disposal orbit 300 km above GEO for end-of-life satellites."],
        ["Hohmann Transfer", "A minimum-energy two-burn transfer between coplanar circular orbits."],
        ["Impulse", "An instantaneous change in velocity (idealised thruster firing)."],
        ["Intent Assessment", "Scoring of adversary intercept intent by correlating observed behaviour with attack profiles and geometric opportunity."],
        ["Intercept Envelope", "The probabilistic reachability volume of a satellite within a given ΔV budget and time horizon."],
        ["J2", "The dominant zonal harmonic of Earth's gravitational field, caused by equatorial bulge."],
        ["J2 Drift", "Exploitation of J2-induced RAAN precession for fuel-free orbital plane alignment."],
        ["Lambert Problem", "The two-point boundary-value problem: find the orbit connecting two positions in a given time."],
        ["Manoeuvre Classification", "Estimating the type and magnitude of a manoeuvre from observed orbital element changes."],
        ["Min-Time Intercept", "The shortest time-of-flight intercept achievable within a given ΔV budget, found by binary search on TOF."],
        ["Miss Distance", "The closest approach distance between two objects at the intercept epoch."],
        ["NMC", "Natural Motion Circumnavigation — passive relative orbit for inspection/proximity operations."],
        ["NORAD ID", "Five-digit catalogue number assigned to each tracked space object."],
        ["Orbital Terrain", "Altitude-vs-inclination risk map combining debris density, congestion, and radiation environment."],
        ["Passive Safety", "Property of a relative orbit where natural motion prevents collision if propulsion fails."],
        ["Perigee", "The point in an orbit closest to the central body."],
        ["Phasing Orbit", "A temporary orbit with a different period, used to close along-track separation over N revolutions."],
        ["Plane Change", "A manoeuvre that alters the orbital inclination or RAAN."],
        ["Prograde", "In the direction of the velocity vector."],
        ["RAAN", "Right Ascension of the Ascending Node — the angle from the vernal equinox to the ascending node."],
        ["Retrograde", "Opposite to the direction of the velocity vector."],
        ["Semi-major Axis", "Half the longest diameter of an elliptical orbit."],
        ["SGP4", "Simplified General Perturbations model 4 — the standard TLE propagator."],
        ["Stability Score", "Ratio of actual Δvy₀ to the CW bounded-motion requirement (Δvy₀ = −2n·Δx₀); 1.0 = bounded relative motion."],
        ["TLE", "Two-Line Element set — the standard format for satellite orbital data."],
        ["Time of Flight", "Duration of the transfer from departure to arrival."],
        ["Trade-Space", "The set of feasible solutions plotted by their key trade-off parameters."],
        ["UDL", "Unified Data Library — authoritative source for TLE data and space object notifications."],
        ["UDL Data Mode", "Classification tag on UDL records: REAL (operational), SIMULATED, EXERCISE, or TEST."],
        ["VNB Frame", "Velocity-Normal-Binormal local orbital reference frame."],
        ["Vis-viva Equation", "Relates orbital velocity to position and semi-major axis."],
        ["GCAT", "General Catalogue — Jonathan McDowell's authoritative space object catalogue at planet4589.org."],
        ["HRR", "High Rate Revisit — JCO satellite watchlist identifying high-priority objects by threat rank."],
        ["Pattern of Life (PoL)", "Analysis of a satellite's full TLE history to characterise manoeuvre behaviour and detect anomalies."],
        ["PoL Status", "Classification of a satellite's historical behaviour: NOMINAL (routine station-keeping), ANOMALOUS (unusual or large manoeuvres), or SUSPICIOUS (evasive or targeting behaviour)."],
        ["TLE Source", "The UDL data provider (e.g. 18 SDS, LeoLabs) whose elsets are used for TLE fetches."],
        ["Drift Phase", "A continuous period of east-west longitude drift at a measurable rate, detected from successive GEO elsets."],
    ]
    _add_data_table(doc, ["Term", "Definition"], terms)


def _section_15_pattern_of_life(doc: Document) -> None:
    """Section 15: Historical Pattern of Life Analysis."""
    _add_page_break(doc)
    _add_eyebrow(doc, "Section 15")
    _add_heading_h1(doc, "Historical Pattern of Life Analysis")

    _add_body(doc,
        "The Historical Pattern of Life (PoL) panel analyses the full TLE history "
        "of a satellite to detect past manoeuvres, characterise station-keeping "
        "behaviour, and identify anomalous activity. It fetches up to 5,000 "
        "historical elsets from UDL and combines them with the latest current elset "
        "to produce a comprehensive orbital element timeline."
    )

    _add_heading_h2(doc, "Operator Workflow")
    _add_bullet(doc, "1. Open the Historical PoL hero tab.")
    _add_bullet(doc, "2. Enter the SATNO to analyse.")
    _add_bullet(doc, "3. Adjust the \u0394V Detection Threshold (default 2.0 m/s) — lower values detect smaller station-keeping manoeuvres; raise it to focus on significant repositioning events.")
    _add_bullet(doc, "4. Choose the data source: UDL session (uses active credentials) or enter UDL credentials directly.")
    _add_bullet(doc, "5. Click Analyse \u2014 SIPC fetches TLE history and the current elset, merges and sorts them chronologically, then runs the PoL engine.")
    _add_bullet(doc, "6. Review the results: manoeuvre timeline, orbital element charts, drift phase analysis, and PoL status.")

    _add_heading_h2(doc, "Results Panels")
    _add_data_table(doc,
        ["Panel", "Description"],
        [
            ["Summary", "Total TLE count, manoeuvre count, PoL status (NOMINAL / ANOMALOUS / SUSPICIOUS)"],
            ["Manoeuvre Timeline", "Table of detected manoeuvres with epoch, estimated \u0394V (m/s), manoeuvre type, and longitude at time of manoeuvre"],
            ["Orbital Elements Chart", "Time-series plots of altitude, inclination, eccentricity, RAAN, and period across the full TLE history"],
            ["GEO Longitude Chart", "Longitude and drift rate over time for GEO/near-GEO objects, with drift phase bands"],
            ["Drift Phases", "Table of distinct longitude drift phases: start/end epoch, direction, drift rate (\u00B0/day), start/end longitude"],
            ["Statistics", "Mean \u0394V, 2\u03C3 high-\u0394V threshold, mean manoeuvre interval, and interval bounds"],
        ],
    )

    _add_heading_h2(doc, "Manoeuvre Classification")
    _add_data_table(doc,
        ["Type", "Description"],
        [
            ["station_keeping", "Small periodic manoeuvre to maintain orbital slot or altitude"],
            ["plane_change", "Change in inclination or RAAN"],
            ["repositioning", "Large manoeuvre indicating a change of operational orbit or longitude slot"],
            ["unknown", "Manoeuvre detected but type cannot be classified from TLE data alone"],
        ],
    )

    _add_callout(doc,
        "PoL analysis requires at least 10 valid TLEs for meaningful results. "
        "Objects with sparse TLE histories will produce limited manoeuvre detection. "
        "The dual-fetch strategy (historical + current) ensures the most recent elset "
        "is always included even when the 5,000-record cap is reached."
    )


def _section_16_gcat_browser(doc: Document) -> None:
    """Section 16: GCAT Browser."""
    _add_page_break(doc)
    _add_eyebrow(doc, "Section 16")
    _add_heading_h1(doc, "GCAT Browser")

    _add_body(doc,
        "The GCAT (General Catalogue) Browser provides instant access to Jonathan "
        "McDowell's authoritative space object catalogue datasets, hosted at "
        "planet4589.org. 28 datasets spanning launch history, payload attributes, "
        "orbital decay records, and deep-space catalogues are available."
    )

    _add_heading_h2(doc, "Accessing the Browser")
    _add_body(doc,
        "Open the GCAT Browser hero tab. The 28 dataset tiles load immediately. "
        "Click any tile to fetch and display that dataset in a searchable, "
        "sortable, paginated table. No UDL connection is required \u2014 data is "
        "fetched directly from planet4589.org on demand."
    )

    _add_heading_h2(doc, "Available Dataset Categories")
    _add_data_table(doc,
        ["Category", "Example Datasets"],
        [
            ["Launch catalogue", "Satcat, Launch log, Launch sites"],
            ["Payload attributes", "Payload catalogue, Owner/operator, Masses"],
            ["Orbital decay", "Decay log, Reentry predictions"],
            ["Deep space", "Deep space catalogue, Lunar missions, Planetary missions"],
            ["Debris & fragmentation", "Fragmentation events, Debris catalogue"],
            ["Auxiliary", "Rocket body catalogue, Infrastructure"],
        ],
    )

    _add_heading_h2(doc, "Search and Navigation")
    _add_body(doc,
        "Each dataset table supports:"
    )
    _add_bullet(doc, "Text search \u2014 filters all visible columns in real time.")
    _add_bullet(doc, "Column sort \u2014 click any header to sort ascending/descending.")
    _add_bullet(doc, "Pagination \u2014 navigate large datasets in pages of 50 rows.")
    _add_callout(doc,
        "GCAT data is fetched on demand and cached for the session. Datasets are "
        "not updated automatically \u2014 click the dataset tile again to refresh."
    )


def _section_17_appendices(doc: Document) -> None:
    """Section 17: Appendices."""
    _add_page_break(doc)
    _add_eyebrow(doc, "Section 17")
    _add_heading_h1(doc, "Appendices")

    _add_heading_h2(doc, "A. Keyboard Shortcuts & UI Tips")
    _add_data_table(doc,
        ["Action", "How"],
        [
            ["Quick-add satellite", "Enter NORAD ID and press Enter"],
            ["Run intercept", "Click Apply Intercept or press Enter in the form"],
            ["View orbital events", "Click Compute Events with satellites selected"],
            ["Clear trade-space", "Click Clear History on the scatter plot"],
            ["Switch TLE mode", "Use the radio buttons (UDL / Manual / Catalogue)"],
        ],
    )

    _add_heading_h2(doc, "B. Common Error Messages")
    _add_data_table(doc,
        ["Error", "Cause", "Resolution"],
        [
            ["No TLE found", "Satellite not in session", "Add the satellite first"],
            ["Intercept calculation failed", "Solver diverged or invalid geometry", "Try a longer time of flight or different coast duration"],
            ["UDL connection failed", "Invalid credentials or network error", "Re-enter UDL credentials; check network connectivity"],
            ["No orbital events found", "TLE epoch too stale or window too short", "Use a fresher TLE or widen the scenario window"],
        ],
    )

    _add_heading_h2(doc, "C. API Endpoints")
    _add_data_table(doc,
        ["Method", "Endpoint", "Purpose"],
        [
            ["POST", "/plan/maneuver/apply-intercept", "Run an intercept calculation"],
            ["POST", "/plan/maneuver/apply-all-intercepts", "Run all intercepts (bulk)"],
            ["GET", "/plan/maneuver/orbital-events", "Compute orbital events for satellites"],
            ["GET", "/plan/maneuver/trade-space-data", "Retrieve trade-space JSON for charting"],
            ["POST", "/plan/maneuver/clear-history", "Clear intercept history and reset chart"],
            ["GET", "/plan/threat/target-config", "Load threat sweep target group dropdown"],
            ["POST", "/plan/threat/fetch-targets", "Pre-fetch TLEs for a target group"],
            ["POST", "/plan/threat/sweep", "Run batch Hohmann + Lambert threat sweep"],
            ["POST", "/plan/threat/refine", "Lambert-refine a single sweep entry"],
            ["GET", "/pol/panel", "Load the Historical PoL panel"],
            ["POST", "/pol/analyse", "Run PoL analysis for a SATNO"],
            ["GET", "/pol/chart-data/{satno}", "Retrieve PoL chart JSON for a SATNO"],
            ["POST", "/assets/blue/quick-add", "One-click HRR → blue asset ingestion"],
            ["POST", "/assets/red/quick-add", "One-click HRR → red track ingestion"],
            ["POST", "/blue/add", "Add a blue asset to the session"],
            ["POST", "/red/add", "Add a red track to the session"],
            ["POST", "/blue/remove", "Remove a blue asset"],
            ["POST", "/red/remove", "Remove a red track"],
            ["POST", "/udl/login", "Connect to UDL with credentials"],
            ["POST", "/udl/logout", "Disconnect from UDL"],
            ["POST", "/udl/data-mode", "Set session UDL data mode"],
            ["POST", "/udl/tle-source", "Set preferred TLE source provider"],
            ["GET", "/udl/tle", "Fetch TLE from UDL by NORAD ID (latest or epoch mode)"],
            ["GET", "/udl/statevector", "Fetch latest Cartesian state vector from UDL"],
            ["GET", "/udl/hrr", "Fetch HRR satellite list from UDL"],
            ["GET", "/udl/catalog/search", "Search cached on-orbit catalogue"],
        ],
    )


# ═══════════════════════════════════════════════════════════════════════════
#  Main
# ═══════════════════════════════════════════════════════════════════════════


def generate() -> Path:
    """Generate the complete SIPC Operator Guide and return the output path."""
    doc = _setup_document()
    _add_accent_bar(doc)
    _add_footer(doc)

    # Title page.
    _add_eyebrow(doc, "Bluestaq Ltd")
    _add_heading_h1(doc, "SIPC Operator Guide")
    _add_body(doc, "Satellite Intercept Planning Console")
    _add_body(doc, "Comprehensive Operator Reference \u2014 v6.0")
    _add_body(doc, "")
    _add_metric_card(doc, "23", "MANOEUVRE METHODS")
    _add_metric_card(doc, "17", "GUIDE SECTIONS")
    _add_metric_card(doc, "28", "GCAT DATASETS")

    # Content sections.
    _section_01_introduction(doc)
    _section_02_system_overview(doc)
    _section_03_getting_started(doc)
    _section_04_dashboard(doc)
    _section_05_assets(doc)
    _section_06_scenario_time(doc)
    _section_07_manoeuvre_theory(doc)
    _section_08_tactical_theory(doc)
    _section_09_threat_sweep(doc)
    _section_10_intercept_calculations(doc)
    _section_11_trade_space(doc)
    _section_12_scenario_planning(doc)
    _section_13_scenarios(doc)
    _section_14_glossary(doc)
    _section_15_pattern_of_life(doc)
    _section_16_gcat_browser(doc)
    _section_17_appendices(doc)

    doc.save(str(OUTPUT_PATH))
    return OUTPUT_PATH


if __name__ == "__main__":
    out = generate()
    print(f"Generated: {out}")
